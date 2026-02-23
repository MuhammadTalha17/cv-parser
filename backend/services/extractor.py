import fitz #pymupdf
from docx import Document
import os
import io
import tempfile
from docling.document_converter import DocumentConverter

_converter = DocumentConverter() #instantiate DocumentConverter, once at module load

def extract_text(file_bytes: str, extension: str) -> str:
    tmp_path = None
    try:
        #Docling requires a file path, not bytes like pymupdf.
        suffix = f".{extension}"
        with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name
        
        result = _converter.convert(tmp_path) #file parsing
        return result.document.export_to_markdown() #clean markdown output

    except Exception as e:
        print(f"Error extracting text from {extension.upper()}: {e}")
        raise
    finally:
        #delete temp file
        if tmp_path and os.path.exists(tmp_path):
            os.remove(tmp_path)

# def extract_text_from_pdf(file_bytes: bytes) -> str:
#     try:
#         text = ""
#         with fitz.open(stream=file_bytes, filetype="pdf") as doc:
#             for page in doc:
#                 text += page.get_text()
#         return text.strip()
#     except Exception as e:
#         print(f"Error extracting text from PDF: {e}")
#         raise

# def extract_text_from_docx(file_bytes: bytes) -> str:
#     try:
#         doc = Document(io.BytesIO(file_bytes))
#         # text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
#         # return text.strip()

#         lines = []

#         ["this is line 1", "this is line 2", "this is line 3"]

#         "this is line 1\nthis is line 2\nthis is line 3"

#         """this is line 1
#         this is line 2
#          """

#         # Extract paragraphs
#         for para in doc.paragraphs:
#             if para.text.strip():
#                 lines.append(para.text.strip())

#         # Extract tables
#         for table in doc.tables:
#             for row in table.rows:
#                 #
#                 seen = []
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text and text not in seen:
#                         seen.append(text)
#                 if seen:
#                     lines.append(" | ".join(seen))
        
#         """
#         this is line 1
#         this is line 2
#         this is line 3
#         """
#         return "\n".join(lines).strip()
        
#     except Exception as e:
#         print(f"Error extracting text from DOCX: {e}")
#         raise

