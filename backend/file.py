from docx import Document
from docling.document_converter import DocumentConverter



def extract_text_from_docx(file: str) -> str:
    try:
        doc = Document(file)

        lines = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())

        # Extract tables
        for table in doc.tables:
            for row in table.rows:
                # Deduplicate merged cells (python-docx repeats them)
                seen = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text and text not in seen:
                        seen.append(text)
                if seen:
                    lines.append(" | ".join(seen))

        return "\n".join(lines).strip()
        
    except Exception as e:
        print(f"Error extracting text from DOCX: {e}")
        raise

if __name__ == "__main__":
    source = "c:\\Users\\talha\\Downloads\\muhammad-talha-cv.pdf"  # document per local path or URL
    converter = DocumentConverter()
    result = converter.convert(source)
    print(result.document.export_to_markdown())  # output: "## Docling Technical Report[...]"
    result.document.save_as_markdown(
                    f"./filename.md"
                )